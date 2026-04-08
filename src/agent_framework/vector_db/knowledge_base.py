"""
知识库管理系统
==============

提供文档上传、解析、向量化、检索等功能。

特性:
- 多格式文档支持（PDF、Word、Markdown、TXT）
- 智能分块
- 向量化存储
- 语义检索
- 知识库管理
"""

import os
import re
import uuid
import hashlib
import requests
import heapq
from collections import Counter, defaultdict, OrderedDict
from datetime import datetime
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
import agent_framework.core.fast_json as json
from agent_framework.vector_db.reranker import LightweightReranker
from agent_framework.vector_db.retrieval_utils import (
    normalize_text as _normalize_text_fn,
    tokenize as _tokenize_fn,
    expand_query_tokens as _expand_query_tokens_fn,
    classify_query,
    get_retrieval_weights,
)
from agent_framework.reasoning.model_serving import get_model_serving_manager
from agent_framework.vector_db.vector_store import create_embedding_model_from_endpoint

# 文档解析
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# 向量数据库
try:
    import chromadb
    from chromadb.config import Settings
    CHROMA_SUPPORT = True
except ImportError:
    CHROMA_SUPPORT = False

# 文本分块
from typing import List


@dataclass
class DocumentChunk:
    """文档块"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    content: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = None
    search_data: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    """文档"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    name: str = ""
    type: str = ""  # pdf, docx, md, txt
    size: int = 0
    content: str = ""
    chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())

    def to_dict(self) -> dict:
        return {
            'id': self.id,
            'name': self.name,
            'type': self.type,
            'size': self.size,
            'chunk_count': len(self.chunks),
            'metadata': self.metadata,
            'created_at': self.created_at
        }


@dataclass
class KnowledgeBase:
    """知识库"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    name: str = ""
    description: str = ""
    documents: List[str] = field(default_factory=list)  # document IDs
    metadata: Dict[str, Any] = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())
    updated_at: str = field(default_factory=lambda: datetime.now().isoformat())

    def to_dict(self) -> dict:
        return {
            'id': self.id,
            'name': self.name,
            'description': self.description,
            'document_count': len(self.documents),
            'metadata': self.metadata,
            'created_at': self.created_at,
            'updated_at': self.updated_at
        }


class DocumentParser:
    """文档解析器"""

    @staticmethod
    def parse_txt(file_path: str) -> str:
        """解析TXT文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    @staticmethod
    def parse_md(file_path: str) -> str:
        """解析Markdown文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析PDF文件"""
        if not PDF_SUPPORT:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

        text = ""
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析Word文件"""
        if not DOCX_SUPPORT:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    @classmethod
    def parse(cls, file_path: str, file_type: str = None) -> str:
        """自动解析文档"""
        if file_type is None:
            file_type = os.path.splitext(file_path)[1].lower().lstrip('.')

        parsers = {
            'txt': cls.parse_txt,
            'md': cls.parse_md,
            'markdown': cls.parse_md,
            'pdf': cls.parse_pdf,
            'docx': cls.parse_docx,
            'doc': cls.parse_docx
        }

        parser = parsers.get(file_type)
        if not parser:
            raise ValueError(f"Unsupported file type: {file_type}")

        return parser(file_path)


class TextSplitter:
    """Text splitter with overlap and boundary-aware splitting."""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    @staticmethod
    def _extract_section_meta(chunk_text: str) -> Dict[str, Any]:
        lines = [line.strip() for line in (chunk_text or "").splitlines() if line.strip()]
        if not lines:
            return {"section_title": "", "is_section_heading": False}

        first_line = re.sub(r"\s+", " ", lines[0]).strip(" -#*\t")
        if not first_line:
            return {"section_title": "", "is_section_heading": False}

        heading_like = bool(
            re.match(r"^(chapter|section|part)\b", first_line, re.IGNORECASE)
            or re.match(r"^第[一二三四五六七八九十0-9]+[章节部分条]", first_line)
            or re.match(r"^[一二三四五六七八九十0-9]+[、.)）．]", first_line)
            or (
                len(first_line) <= 40
                and "\n" in (chunk_text or "")[: max(80, len(first_line) + 4)]
                and not re.search(r"[。！？.!?:：；;，,]$", first_line)
            )
        )

        if heading_like or len(first_line) <= 60:
            return {
                "section_title": first_line[:80],
                "is_section_heading": heading_like,
            }
        return {"section_title": "", "is_section_heading": False}

    def split(self, text: str, metadata: Dict[str, Any] = None) -> List[DocumentChunk]:
        """Split text into overlapping chunks."""
        if metadata is None:
            metadata = {}

        chunks: List[DocumentChunk] = []
        start = 0
        text_length = len(text)
        boundaries = ['\n\n', '\n', '\u3002', '\uff01', '\uff1f', '.', '!', '?', '\uff1b', ';', '\uff0c', ',', ' ']

        while start < text_length:
            end = min(start + self.chunk_size, text_length)

            if end < text_length:
                for delimiter in boundaries:
                    pos = text.rfind(delimiter, start, end)
                    if pos != -1 and pos > start + max(20, self.chunk_size // 3):
                        end = pos + len(delimiter)
                        break

            chunk_text = text[start:end].strip()
            if chunk_text:
                section_meta = self._extract_section_meta(chunk_text)
                chunk = DocumentChunk(
                    content=chunk_text,
                    metadata={
                        **metadata,
                        'start': start,
                        'end': end,
                        **({'section_title': section_meta['section_title']} if section_meta.get('section_title') else {}),
                        **({'is_section_heading': section_meta['is_section_heading']} if section_meta.get('section_title') else {}),
                    }
                )
                chunks.append(chunk)

            if end >= text_length:
                break

            next_start = max(start + 1, end - self.chunk_overlap)
            if next_start <= start:
                next_start = end
            start = next_start

        total_chunks = len(chunks)
        for idx, chunk in enumerate(chunks):
            chunk.metadata.setdefault('chunk_index', idx)
            chunk.metadata.setdefault('total_chunks', total_chunks)

        return chunks


class SimpleEmbedding:
    """简单的嵌入生成器（使用哈希，实际应用中应使用真实的embedding模型）"""

    @staticmethod
    def embed(text: str) -> List[float]:
        """生成文本嵌入向量"""
        # 这是一个简化实现，实际应该使用 OpenAI embeddings 或本地模型
        # 这里使用哈希生成固定长度的向量
        hash_obj = hashlib.sha256(text.encode())
        hash_bytes = hash_obj.digest()

        # 转换为384维向量（模拟sentence-transformers的输出）
        vector = []
        for i in range(0, len(hash_bytes), 2):
            if i + 1 < len(hash_bytes):
                value = (hash_bytes[i] * 256 + hash_bytes[i + 1]) / 65535.0
                vector.append(value)

        # 填充到384维
        while len(vector) < 384:
            vector.append(0.0)

        return vector[:384]


class VectorStore:
    """向量存储"""

    def __init__(self, persist_directory: str = "./data/chroma"):
        self.persist_directory = persist_directory
        os.makedirs(persist_directory, exist_ok=True)

        if CHROMA_SUPPORT:
            self.client = chromadb.PersistentClient(path=persist_directory)
        else:
            # 简单的内存存储
            self.collections = {}

    def get_or_create_collection(self, name: str):
        """获取或创建集合"""
        if CHROMA_SUPPORT:
            return self.client.get_or_create_collection(name=name)
        else:
            if name not in self.collections:
                self.collections[name] = {
                    'documents': [],
                    'embeddings': [],
                    'metadatas': [],
                    'ids': []
                }
            return self.collections[name]

    def add_chunks(self, collection_name: str, chunks: List[DocumentChunk]):
        """添加文档块"""
        collection = self.get_or_create_collection(collection_name)

        documents = [chunk.content for chunk in chunks]
        embeddings = [chunk.embedding or SimpleEmbedding.embed(chunk.content) for chunk in chunks]
        metadatas = [chunk.metadata for chunk in chunks]
        ids = [chunk.id for chunk in chunks]

        if CHROMA_SUPPORT:
            collection.add(
                documents=documents,
                embeddings=embeddings,
                metadatas=metadatas,
                ids=ids
            )
        else:
            collection['documents'].extend(documents)
            collection['embeddings'].extend(embeddings)
            collection['metadatas'].extend(metadatas)
            collection['ids'].extend(ids)

    def search(
        self,
        collection_name: str,
        query: str,
        top_k: int = 5,
        query_embedding: Optional[List[float]] = None,
    ) -> List[Dict[str, Any]]:
        """搜索相似文档"""
        collection = self.get_or_create_collection(collection_name)
        query_embedding = query_embedding or SimpleEmbedding.embed(query)

        if CHROMA_SUPPORT:
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k
            )

            return [
                {
                    'id': results['ids'][0][i],
                    'content': results['documents'][0][i],
                    'metadata': results['metadatas'][0][i],
                    'distance': results['distances'][0][i] if 'distances' in results else 0
                }
                for i in range(len(results['ids'][0]))
            ]
        else:
            # 简单的余弦相似度计算
            import math

            def cosine_similarity(v1, v2):
                dot_product = sum(a * b for a, b in zip(v1, v2))
                magnitude1 = math.sqrt(sum(a * a for a in v1))
                magnitude2 = math.sqrt(sum(b * b for b in v2))
                if magnitude1 == 0 or magnitude2 == 0:
                    return 0
                return dot_product / (magnitude1 * magnitude2)

            similarities = []
            for i, embedding in enumerate(collection['embeddings']):
                similarity = cosine_similarity(query_embedding, embedding)
                similarities.append((i, similarity))

            # 排序并返回top_k
            similarities.sort(key=lambda x: x[1], reverse=True)

            return [
                {
                    'id': collection['ids'][idx],
                    'content': collection['documents'][idx],
                    'metadata': collection['metadatas'][idx],
                    'distance': 1 - sim
                }
                for idx, sim in similarities[:top_k]
            ]

    def delete_collection(self, name: str):
        """删除集合"""
        if CHROMA_SUPPORT:
            self.client.delete_collection(name=name)
        else:
            if name in self.collections:
                del self.collections[name]


class KnowledgeBaseManager:
    """知识库管理器"""

    DEFAULT_RAG_SETTINGS = {
        "chunk_size": 500,
        "chunk_overlap": 50,
        "search_top_k": 5,
        "candidate_top_k": 30,
        "retrieval_score_threshold": 0.1,
        "mmr_lambda": 0.78,
        "window_size": 1,
        "window_max_chars": 1600,
        "section_max_chars": 2400,
        "query_expansion_enabled": True,
    }

    def __init__(self, data_dir: str = "./data/knowledge"):
        self.data_dir = data_dir
        self.upload_dir = os.path.join(data_dir, "uploads")
        self.kb_dir = os.path.join(data_dir, "knowledge_bases")
        self.doc_dir = os.path.join(data_dir, "documents")

        os.makedirs(self.upload_dir, exist_ok=True)
        os.makedirs(self.kb_dir, exist_ok=True)
        os.makedirs(self.doc_dir, exist_ok=True)

        self.vector_store = VectorStore()
        self.text_splitter = TextSplitter(chunk_size=500, chunk_overlap=50)
        self._kb_chunk_cache: Dict[str, List[DocumentChunk]] = {}
        self._kb_search_cache: Dict[str, Dict[str, Any]] = {}
        self._query_result_cache: Dict[str, OrderedDict] = {}
        self._query_result_cache_limit = 128
        self.reranker = LightweightReranker()

    @classmethod
    def _sanitize_rag_settings(cls, rag_settings: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        raw = dict(rag_settings or {})

        def _int_value(key: str, minimum: int, maximum: int) -> int:
            value = raw.get(key, cls.DEFAULT_RAG_SETTINGS[key])
            try:
                return max(minimum, min(maximum, int(value)))
            except (TypeError, ValueError):
                return int(cls.DEFAULT_RAG_SETTINGS[key])

        def _float_value(key: str, minimum: float, maximum: float) -> float:
            value = raw.get(key, cls.DEFAULT_RAG_SETTINGS[key])
            try:
                return max(minimum, min(maximum, float(value)))
            except (TypeError, ValueError):
                return float(cls.DEFAULT_RAG_SETTINGS[key])

        chunk_size = _int_value("chunk_size", 100, 4000)
        chunk_overlap = _int_value("chunk_overlap", 0, min(1000, max(0, chunk_size - 1)))

        return {
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
            "search_top_k": _int_value("search_top_k", 1, 20),
            "candidate_top_k": _int_value("candidate_top_k", 8, 100),
            "retrieval_score_threshold": round(_float_value("retrieval_score_threshold", 0.0, 1.0), 4),
            "mmr_lambda": round(_float_value("mmr_lambda", 0.0, 1.0), 4),
            "window_size": _int_value("window_size", 0, 4),
            "window_max_chars": _int_value("window_max_chars", 200, 6000),
            "section_max_chars": _int_value("section_max_chars", 400, 8000),
            "query_expansion_enabled": bool(raw.get("query_expansion_enabled", cls.DEFAULT_RAG_SETTINGS["query_expansion_enabled"])),
        }

    def _get_rag_settings(self, kb: Optional[KnowledgeBase]) -> Dict[str, Any]:
        metadata = dict((kb.metadata or {}) if kb else {})
        return self._sanitize_rag_settings(metadata.get("rag_settings") or {})

    def _get_text_splitter_for_kb(self, kb: Optional[KnowledgeBase]) -> TextSplitter:
        rag_settings = self._get_rag_settings(kb)
        return TextSplitter(
            chunk_size=int(rag_settings["chunk_size"]),
            chunk_overlap=int(rag_settings["chunk_overlap"]),
        )

    def _rebuild_document_chunks(self, kb: Optional[KnowledgeBase], doc: Document) -> List[DocumentChunk]:
        splitter = self._get_text_splitter_for_kb(kb)
        source_text = str(doc.content or "").strip()
        if (not source_text or len(source_text) <= 1000) and doc.chunks:
            source_text = self._merge_chunk_texts([chunk.content for chunk in doc.chunks])
        doc.content = source_text
        doc.chunks = splitter.split(
            source_text,
            metadata={"doc_id": doc.id, "doc_name": doc.name},
        )
        for chunk in doc.chunks:
            self._prepare_chunk_search_data(chunk)
        return doc.chunks

    def create_knowledge_base(self, name: str, description: str = "") -> KnowledgeBase:
        """创建知识库"""
        kb = KnowledgeBase(
            name=name,
            description=description,
            metadata={"rag_settings": dict(self.DEFAULT_RAG_SETTINGS)},
        )
        self._save_knowledge_base(kb)
        self._invalidate_kb_cache(kb.id)
        return kb

    def update_knowledge_base_settings(
        self,
        kb_id: str,
        *,
        embedding_endpoint_id: str = "",
        rerank_endpoint_id: str = "",
        rag_settings: Optional[Dict[str, Any]] = None,
    ) -> KnowledgeBase:
        """Update endpoint bindings and advanced RAG settings for a knowledge base."""
        kb = self.get_knowledge_base(kb_id)
        if not kb:
            raise ValueError(f"Knowledge base not found: {kb_id}")

        metadata = dict(kb.metadata or {})
        serving_manager = get_model_serving_manager()

        def _apply_endpoint(endpoint_type: str, endpoint_id: str):
            endpoint_id = (endpoint_id or "").strip()
            keys = [
                f"{endpoint_type}_endpoint_id",
                f"{endpoint_type}_endpoint_name",
                f"{endpoint_type}_backend",
                f"{endpoint_type}_base_url",
            ]
            if not endpoint_id:
                for key in keys:
                    metadata.pop(key, None)
                return

            endpoint = serving_manager.get_endpoint(endpoint_id)
            if endpoint is None:
                raise ValueError(f"{endpoint_type} endpoint not found: {endpoint_id}")
            if endpoint.status != "running":
                raise ValueError(f"{endpoint_type} endpoint is not running: {endpoint_id}")
            if getattr(endpoint, "endpoint_type", "chat") != endpoint_type:
                raise ValueError(f"Endpoint {endpoint_id} is not a {endpoint_type} endpoint")

            metadata[f"{endpoint_type}_endpoint_id"] = endpoint.endpoint_id
            metadata[f"{endpoint_type}_endpoint_name"] = endpoint.model_name
            metadata[f"{endpoint_type}_backend"] = endpoint.backend
            metadata[f"{endpoint_type}_base_url"] = endpoint.base_url

        _apply_endpoint("embedding", embedding_endpoint_id)
        _apply_endpoint("rerank", rerank_endpoint_id)
        metadata["rag_settings"] = self._sanitize_rag_settings(rag_settings or metadata.get("rag_settings") or {})

        kb.metadata = metadata
        kb.updated_at = datetime.now().isoformat()
        self._save_knowledge_base(kb)
        self._invalidate_kb_cache(kb.id)
        return kb

    def rebuild_knowledge_base_vectors(self, kb_id: str) -> KnowledgeBase:
        """Rebuild vector embeddings for all documents in the knowledge base."""
        kb = self.get_knowledge_base(kb_id)
        if not kb:
            raise ValueError(f"Knowledge base not found: {kb_id}")

        chunks: List[DocumentChunk] = []
        for doc_id in kb.documents:
            doc = self.get_document(doc_id)
            if not doc:
                continue
            chunks.extend(self._rebuild_document_chunks(kb, doc))
            self._save_document(doc)

        self.vector_store.delete_collection(kb_id)

        if chunks:
            self._embed_chunks_for_kb(kb, chunks)
            self.vector_store.add_chunks(kb_id, chunks)

        kb.updated_at = datetime.now().isoformat()
        self._save_knowledge_base(kb)
        self._invalidate_kb_cache(kb_id)
        return kb

    def get_knowledge_base(self, kb_id: str) -> Optional[KnowledgeBase]:
        """获取知识库"""
        kb_file = os.path.join(self.kb_dir, f"{kb_id}.json")
        if not os.path.exists(kb_file):
            return None

        with open(kb_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            return KnowledgeBase(**data)

    def list_knowledge_bases(self) -> List[KnowledgeBase]:
        """列出所有知识库"""
        kbs = []
        for filename in os.listdir(self.kb_dir):
            if filename.endswith('.json'):
                kb_id = filename[:-5]
                kb = self.get_knowledge_base(kb_id)
                if kb:
                    kbs.append(kb)
        return kbs

    def delete_knowledge_base(self, kb_id: str):
        """删除知识库"""
        kb = self.get_knowledge_base(kb_id)
        if not kb:
            return

        # 删除向量集合
        self.vector_store.delete_collection(kb_id)

        # 删除知识库文件
        kb_file = os.path.join(self.kb_dir, f"{kb_id}.json")
        if os.path.exists(kb_file):
            os.remove(kb_file)
        self._invalidate_kb_cache(kb_id)

    def upload_document(self, kb_id: str, file_path: str, filename: str) -> Document:
        """上传文档到知识库"""
        kb = self.get_knowledge_base(kb_id)
        if not kb:
            raise ValueError(f"Knowledge base not found: {kb_id}")

        # 解析文档
        file_type = os.path.splitext(filename)[1].lower().lstrip('.')
        content = DocumentParser.parse(file_path, file_type)

        # 创建文档对象
        doc = Document(
            name=filename,
            type=file_type,
            size=os.path.getsize(file_path),
            content=content,
            metadata={'kb_id': kb_id}
        )

        # 分块
        self._rebuild_document_chunks(kb, doc)

        self._embed_chunks_for_kb(kb, doc.chunks)

        # 向量化并存储
        self.vector_store.add_chunks(kb_id, doc.chunks)

        # 保存文档
        self._save_document(doc)

        # 更新知识库
        kb.documents.append(doc.id)
        kb.updated_at = datetime.now().isoformat()
        self._save_knowledge_base(kb)
        self._invalidate_kb_cache(kb_id)

        return doc

    def get_document(self, doc_id: str) -> Optional[Document]:
        """获取文档"""
        doc_file = os.path.join(self.doc_dir, f"{doc_id}.json")
        if not os.path.exists(doc_file):
            return None

        with open(doc_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            # 重建chunks
            chunks = [DocumentChunk(**chunk) for chunk in data.pop('chunks', [])]
            doc = Document(**data)
            doc.chunks = chunks
            return doc

    def delete_document(self, kb_id: str, doc_id: str):
        """Delete a document from the knowledge base."""
        kb = self.get_knowledge_base(kb_id)
        if not kb:
            return

        if doc_id in kb.documents:
            kb.documents.remove(doc_id)
            kb.updated_at = datetime.now().isoformat()
            self._save_knowledge_base(kb)

        doc_file = os.path.join(self.doc_dir, f"{doc_id}.json")
        if os.path.exists(doc_file):
            os.remove(doc_file)
        self._invalidate_kb_cache(kb_id)

    def _invalidate_kb_cache(self, kb_id: str):
        self._kb_chunk_cache.pop(kb_id, None)
        self._kb_search_cache.pop(kb_id, None)
        self._query_result_cache.pop(kb_id, None)

    @staticmethod
    def _bounded_score(value: Any) -> float:
        try:
            return max(0.0, min(1.0, float(value)))
        except (TypeError, ValueError):
            return 0.0

    def _get_endpoint_by_id(self, endpoint_id: str, endpoint_type: str):
        endpoint_id = str(endpoint_id or "").strip()
        if not endpoint_id:
            return None
        endpoint = get_model_serving_manager().get_endpoint(endpoint_id)
        if endpoint is None:
            raise ValueError(f"{endpoint_type} endpoint not found: {endpoint_id}")
        if endpoint.status != "running":
            raise ValueError(f"{endpoint_type} endpoint is not running: {endpoint_id}")
        if getattr(endpoint, "endpoint_type", "chat") != endpoint_type:
            raise ValueError(f"Endpoint {endpoint_id} is not a {endpoint_type} endpoint")
        return endpoint

    def _get_configured_endpoint(self, kb: Optional[KnowledgeBase], endpoint_type: str):
        mgr = get_model_serving_manager()
        if kb is None:
            return mgr.get_best_endpoint(endpoint_type)
        endpoint_id = str((kb.metadata or {}).get(f"{endpoint_type}_endpoint_id", "") or "").strip()
        if endpoint_id:
            try:
                return self._get_endpoint_by_id(endpoint_id, endpoint_type)
            except ValueError:
                pass
        return mgr.get_best_endpoint(endpoint_type)

    def _resolve_endpoint(
        self,
        kb: Optional[KnowledgeBase],
        endpoint_type: str,
        override_endpoint_id: str = "",
    ):
        override_endpoint_id = str(override_endpoint_id or "").strip()
        if override_endpoint_id:
            return self._get_endpoint_by_id(override_endpoint_id, endpoint_type)
        return self._get_configured_endpoint(kb, endpoint_type)

    def _embed_chunks_for_kb(self, kb: Optional[KnowledgeBase], chunks: List[DocumentChunk]):
        endpoint = self._get_configured_endpoint(kb, "embedding")
        if endpoint is None or not chunks:
            return

        embedder = create_embedding_model_from_endpoint(endpoint.endpoint_id)
        batch_size = 16
        for start in range(0, len(chunks), batch_size):
            batch = chunks[start:start + batch_size]
            embeddings = embedder.embed_batch([chunk.content for chunk in batch])
            for chunk, embedding in zip(batch, embeddings):
                chunk.embedding = embedding

    def _embed_query_for_kb(
        self,
        kb: Optional[KnowledgeBase],
        query: str,
        embedding_endpoint_id: str = "",
    ) -> Optional[List[float]]:
        endpoint = self._resolve_endpoint(kb, "embedding", embedding_endpoint_id)
        if endpoint is None:
            return None
        embedder = create_embedding_model_from_endpoint(endpoint.endpoint_id)
        return embedder.embed_text(query)

    @staticmethod
    def _parse_remote_rerank_results(payload: Any, expected_count: int) -> Optional[List[Dict[str, Any]]]:
        if not isinstance(payload, dict):
            return None

        results = payload.get("results") or payload.get("data") or payload.get("items")
        if not isinstance(results, list):
            return None

        parsed = [
            {"remote_score": 0.0, "source": "remote_rerank"}
            for _ in range(expected_count)
        ]

        for rank, item in enumerate(results):
            if not isinstance(item, dict):
                continue
            try:
                index = int(item.get("index", rank))
            except (TypeError, ValueError):
                index = rank
            if index < 0 or index >= expected_count:
                continue
            score = (
                item.get("relevance_score")
                if item.get("relevance_score") is not None
                else item.get("score")
            )
            if score is None:
                score = item.get("similarity")
            parsed[index] = {
                "remote_score": max(0.0, float(score or 0.0)),
                "source": "remote_rerank",
                "raw": item,
            }

        return parsed

    def _remote_rerank_documents(
        self,
        kb: Optional[KnowledgeBase],
        query: str,
        documents: List[str],
        rerank_endpoint_id: str = "",
    ) -> Optional[List[Dict[str, Any]]]:
        endpoint = self._resolve_endpoint(kb, "rerank", rerank_endpoint_id)
        if endpoint is None or not documents:
            return None

        headers = {"Content-Type": "application/json"}
        if endpoint.api_key:
            headers["Authorization"] = f"Bearer {endpoint.api_key}"

        base_url = str(endpoint.base_url or "").rstrip("/")
        candidates = [f"{base_url}/rerank"]
        if base_url.endswith("/v1"):
            candidates.append(f"{base_url[:-3]}/rerank")

        payload = {
            "model": endpoint.model_name,
            "query": query,
            "documents": documents,
            "top_n": len(documents),
            "return_documents": False,
        }

        last_error: Optional[Exception] = None
        for url in dict.fromkeys(candidates):
            try:
                response = requests.post(url, headers=headers, json=payload, timeout=30)
                response.raise_for_status()
                parsed = self._parse_remote_rerank_results(response.json(), len(documents))
                if parsed is not None:
                    return parsed
            except Exception as exc:
                last_error = exc
                continue

        if last_error:
            print(f"[knowledge_base] remote rerank fallback: {last_error}")
        return None

    @staticmethod
    def _normalize_text(text: str) -> str:
        return _normalize_text_fn(text)

    @classmethod
    def _expand_query_tokens(cls, query: str) -> List[str]:
        return list(_expand_query_tokens_fn(query))

    @classmethod
    def _tokenize_for_search(cls, text: str) -> List[str]:
        return list(_tokenize_fn(text))

    def _build_query_profile(self, query: str, *, query_expansion_enabled: bool = True) -> Dict[str, Any]:
        query_tokens = tuple(
            self._expand_query_tokens(query)
            if query_expansion_enabled
            else self._tokenize_for_search(query)
        )
        query_counter = Counter(query_tokens)
        return {
            'query': query,
            'tokens': query_tokens,
            'counter': query_counter,
            'normalized_query': self._normalize_text(query),
            'query_terms': {
                token for token in query_counter
                if len(token) > 1 or re.search(r"[\u4e00-\u9fff]", token)
            },
            'total': max(1, sum(query_counter.values())),
        }

    def _prepare_chunk_search_data(self, chunk: DocumentChunk) -> Dict[str, Any]:
        if chunk.search_data:
            tokens = chunk.search_data.get('tokens') or []
            normalized_content = chunk.search_data.get('normalized_content', '')
            if tokens and normalized_content:
                return chunk.search_data

        metadata = dict(chunk.metadata or {})
        doc_name = str(metadata.get('doc_name', '') or '')
        section_title = str(metadata.get('section_title', '') or '')
        tokens = list(self._tokenize_for_search(chunk.content or ''))
        chunk.search_data = {
            'tokens': tokens,
            'normalized_content': self._normalize_text(chunk.content or ''),
            'normalized_doc_name': self._normalize_text(doc_name),
            'normalized_section_title': self._normalize_text(section_title),
        }
        return chunk.search_data

    @staticmethod
    def _clone_search_results(results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        cloned: List[Dict[str, Any]] = []
        for item in results:
            cloned.append({
                **item,
                'metadata': dict(item.get('metadata') or {}),
            })
        return cloned

    def _get_cached_query_results(
        self,
        kb_id: str,
        query: str,
        top_k: int,
        embedding_endpoint_id: str = "",
        rerank_endpoint_id: str = "",
    ) -> Optional[List[Dict[str, Any]]]:
        kb_cache = self._query_result_cache.get(kb_id)
        if not kb_cache:
            return None

        cache_key = (
            self._normalize_text(query),
            int(top_k),
            str(embedding_endpoint_id or "").strip(),
            str(rerank_endpoint_id or "").strip(),
        )
        cached = kb_cache.get(cache_key)
        if cached is None:
            return None
        kb_cache.move_to_end(cache_key)
        return self._clone_search_results(cached)

    def _cache_query_results(
        self,
        kb_id: str,
        query: str,
        top_k: int,
        results: List[Dict[str, Any]],
        embedding_endpoint_id: str = "",
        rerank_endpoint_id: str = "",
    ) -> None:
        kb_cache = self._query_result_cache.setdefault(kb_id, OrderedDict())
        cache_key = (
            self._normalize_text(query),
            int(top_k),
            str(embedding_endpoint_id or "").strip(),
            str(rerank_endpoint_id or "").strip(),
        )
        kb_cache[cache_key] = self._clone_search_results(results)
        kb_cache.move_to_end(cache_key)
        while len(kb_cache) > self._query_result_cache_limit:
            kb_cache.popitem(last=False)

    def _get_kb_chunks(self, kb_id: str) -> List[DocumentChunk]:
        cached = self._kb_chunk_cache.get(kb_id)
        if cached is not None:
            return cached

        kb = self.get_knowledge_base(kb_id)
        if not kb:
            return []

        chunks: List[DocumentChunk] = []
        for doc_id in kb.documents:
            doc = self.get_document(doc_id)
            if not doc:
                continue
            for idx, chunk in enumerate(doc.chunks):
                chunk.metadata.setdefault('doc_id', doc.id)
                chunk.metadata.setdefault('doc_name', doc.name)
                chunk.metadata.setdefault('chunk_index', idx)
                chunk.metadata.setdefault('total_chunks', len(doc.chunks))
                self._prepare_chunk_search_data(chunk)
                chunks.append(chunk)

        self._kb_chunk_cache[kb_id] = chunks
        return chunks

    def _get_kb_search_data(self, kb_id: str) -> Dict[str, Any]:
        cached = self._kb_search_cache.get(kb_id)
        if cached is not None:
            return cached

        chunks = self._get_kb_chunks(kb_id)
        by_id: Dict[str, DocumentChunk] = {}
        by_doc: Dict[str, List[DocumentChunk]] = defaultdict(list)
        search_entries: List[Dict[str, Any]] = []

        for chunk in chunks:
            metadata = dict(chunk.metadata or {})
            doc_id = str(metadata.get('doc_id', '') or '')
            doc_name = str(metadata.get('doc_name', '') or '')
            section_title = str(metadata.get('section_title', '') or '')
            chunk_index = int(metadata.get('chunk_index', 0) or 0)
            chunk_search_data = self._prepare_chunk_search_data(chunk)
            normalized_content = str(chunk_search_data.get('normalized_content', '') or '')
            tokens = tuple(chunk_search_data.get('tokens') or ())

            entry = {
                'chunk': chunk,
                'id': chunk.id,
                'doc_id': doc_id,
                'doc_name': doc_name,
                'section_title': section_title,
                'chunk_index': chunk_index,
                'content': chunk.content or '',
                'normalized_content': normalized_content,
                'tokens': tokens,
                'token_set': set(tokens),
                'token_counter': Counter(tokens),
                'normalized_doc_name': str(chunk_search_data.get('normalized_doc_name', '') or self._normalize_text(doc_name)),
                'normalized_section_title': str(chunk_search_data.get('normalized_section_title', '') or self._normalize_text(section_title)),
                'is_section_heading': bool(metadata.get('is_section_heading') is True or self._is_section_heading_chunk(chunk)),
                'metadata': metadata,
            }
            by_id[chunk.id] = chunk
            if doc_id:
                by_doc[doc_id].append(chunk)
            search_entries.append(entry)

        heading_indices_by_doc: Dict[str, List[int]] = {}
        for doc_id, doc_chunks in by_doc.items():
            doc_chunks.sort(key=lambda item: int((item.metadata or {}).get('chunk_index', 0) or 0))
            heading_indices_by_doc[doc_id] = [
                int((chunk.metadata or {}).get('chunk_index', 0) or 0)
                for chunk in doc_chunks
                if bool((chunk.metadata or {}).get('is_section_heading') is True or self._is_section_heading_chunk(chunk))
            ]

        # Build inverted index: token → set of entry indices
        inverted_index: Dict[str, set] = defaultdict(set)
        for entry_idx, entry in enumerate(search_entries):
            for token in entry['token_set']:
                inverted_index[token].add(entry_idx)

        cached = {
            'chunks': chunks,
            'by_id': by_id,
            'by_doc': dict(by_doc),
            'search_entries': search_entries,
            'entry_map': {entry['id']: entry for entry in search_entries},
            'inverted_index': dict(inverted_index),
            'heading_indices_by_doc': heading_indices_by_doc,
            'chunk_window_cache': {},
            'section_window_cache': {},
        }
        self._kb_search_cache[kb_id] = cached
        return cached

    def _build_chunk_window(
        self,
        kb_id: str,
        chunk_id: str,
        *,
        window: int = 1,
        max_chars: int = 1600,
    ) -> Dict[str, Any]:
        search_data = self._get_kb_search_data(kb_id)
        cache_key = (chunk_id, window, max_chars)
        cached = search_data['chunk_window_cache'].get(cache_key)
        if cached is not None:
            return cached

        target = search_data['by_id'].get(chunk_id)
        if not target:
            return {
                'content': '',
                'chunk_ids': [],
                'window_start': 0,
                'window_end': 0,
            }

        doc_id = str((target.metadata or {}).get('doc_id', ''))
        if not doc_id:
            result = {
                'content': target.content,
                'chunk_ids': [chunk_id],
                'window_start': int((target.metadata or {}).get('chunk_index', 0) or 0),
                'window_end': int((target.metadata or {}).get('chunk_index', 0) or 0),
            }
            search_data['chunk_window_cache'][cache_key] = result
            return result

        doc_chunks = list(search_data['by_doc'].get(doc_id, []))

        target_idx = int((target.metadata or {}).get('chunk_index', 0) or 0)
        start_idx = max(0, target_idx - window)
        end_idx = min(len(doc_chunks) - 1, target_idx + window)

        selected = doc_chunks[start_idx:end_idx + 1]
        merged_parts: List[str] = []
        merged_ids: List[str] = []
        total_chars = 0

        for chunk in selected:
            piece = (chunk.content or '').strip()
            if not piece:
                continue
            projected = total_chars + len(piece) + (2 if merged_parts else 0)
            if merged_parts and projected > max_chars:
                break
            merged_parts.append(piece)
            merged_ids.append(chunk.id)
            total_chars = projected

        if not merged_parts:
            merged_parts = [(target.content or '').strip()]
            merged_ids = [target.id]
            start_idx = end_idx = target_idx

        result = {
            'content': self._merge_chunk_texts(merged_parts),
            'chunk_ids': merged_ids,
            'window_start': start_idx,
            'window_end': start_idx + max(0, len(merged_ids) - 1),
        }
        search_data['chunk_window_cache'][cache_key] = result
        return result

    @staticmethod
    def _trim_leading_overlap(previous: str, current: str, min_overlap: int = 18, max_overlap: int = 140) -> str:
        previous = previous or ''
        current = current or ''
        if not previous or not current:
            return current

        upper = min(len(previous), len(current), max_overlap)
        for size in range(upper, min_overlap - 1, -1):
            if previous[-size:] == current[:size]:
                return current[size:]
        return current

    def _merge_chunk_texts(self, parts: List[str]) -> str:
        cleaned_parts = [str(part or '').strip() for part in parts if str(part or '').strip()]
        if not cleaned_parts:
            return ''

        merged = cleaned_parts[0]
        for piece in cleaned_parts[1:]:
            trimmed_piece = self._trim_leading_overlap(merged, piece)
            if not trimmed_piece:
                continue
            if merged.endswith(('\n', ' ', '。', '.', '!', '?', '！', '？')):
                separator = '' if trimmed_piece.startswith(('\n', ' ', '。', '.', '!', '?', '！', '？')) else ' '
                merged += separator + trimmed_piece
            else:
                merged += "\n\n" + trimmed_piece
        return merged

    def _is_section_heading_chunk(self, chunk: DocumentChunk) -> bool:
        metadata = chunk.metadata or {}
        if metadata.get('is_section_heading') is True:
            return True

        lines = [line.strip() for line in (chunk.content or '').splitlines() if line.strip()]
        if not lines:
            return False

        first_line = re.sub(r"\s+", " ", lines[0]).strip(" -#*\t")
        if not first_line:
            return False

        return bool(
            re.match(r"^(chapter|section|part)\b", first_line, re.IGNORECASE)
            or re.match(r"^第[一二三四五六七八九十0-9]+[章节部分条]", first_line)
            or re.match(r"^[一二三四五六七八九十0-9]+[、.)）．]", first_line)
            or (
                len(first_line) <= 40
                and '\n' in (chunk.content or '')[: max(80, len(first_line) + 4)]
                and not re.search(r"[。！？.!?:：；;，,]$", first_line)
            )
        )

    def _build_section_window(
        self,
        kb_id: str,
        chunk_id: str,
        *,
        max_chars: int = 2400,
    ) -> Dict[str, Any]:
        search_data = self._get_kb_search_data(kb_id)
        cache_key = (chunk_id, max_chars)
        cached = search_data['section_window_cache'].get(cache_key)
        if cached is not None:
            return cached

        target = search_data['by_id'].get(chunk_id)
        if not target:
            return {
                'content': '',
                'chunk_ids': [],
                'section_start': 0,
                'section_end': 0,
                'section_title': '',
            }

        doc_id = str((target.metadata or {}).get('doc_id', ''))
        if not doc_id:
            index = int((target.metadata or {}).get('chunk_index', 0) or 0)
            result = {
                'content': target.content,
                'chunk_ids': [target.id],
                'section_start': index,
                'section_end': index,
                'section_title': str((target.metadata or {}).get('section_title', '') or ''),
            }
            search_data['section_window_cache'][cache_key] = result
            return result

        doc_chunks = list(search_data['by_doc'].get(doc_id, []))

        target_idx = int((target.metadata or {}).get('chunk_index', 0) or 0)
        heading_indices = list(search_data['heading_indices_by_doc'].get(doc_id, []))

        if not heading_indices:
            result = {
                'content': target.content,
                'chunk_ids': [target.id],
                'section_start': target_idx,
                'section_end': target_idx,
                'section_title': str((target.metadata or {}).get('section_title', '') or ''),
            }
            search_data['section_window_cache'][cache_key] = result
            return result

        section_start = max((idx for idx in heading_indices if idx <= target_idx), default=heading_indices[0])
        following_headings = [idx for idx in heading_indices if idx > target_idx]
        section_end = following_headings[0] - 1 if following_headings else len(doc_chunks) - 1
        section_start = max(0, min(section_start, len(doc_chunks) - 1))
        section_end = max(section_start, min(section_end, len(doc_chunks) - 1))

        selected = doc_chunks[section_start:section_end + 1]
        merged_parts: List[str] = []
        merged_ids: List[str] = []
        total_chars = 0

        for chunk in selected:
            piece = (chunk.content or '').strip()
            if not piece:
                continue
            projected = total_chars + len(piece) + (2 if merged_parts else 0)
            if merged_parts and projected > max_chars:
                break
            merged_parts.append(piece)
            merged_ids.append(chunk.id)
            total_chars = projected

        if not merged_parts:
            merged_parts = [(target.content or '').strip()]
            merged_ids = [target.id]
            section_start = section_end = target_idx

        section_title = ''
        if selected:
            section_title = str((selected[0].metadata or {}).get('section_title', '') or '')
        if not section_title:
            section_title = str((target.metadata or {}).get('section_title', '') or '')

        result = {
            'content': self._merge_chunk_texts(merged_parts),
            'chunk_ids': merged_ids,
            'section_start': section_start,
            'section_end': section_start + max(0, len(merged_ids) - 1),
            'section_title': section_title,
        }
        search_data['section_window_cache'][cache_key] = result
        return result

    def _lexical_score(
        self,
        query: str,
        content: str,
        metadata: Optional[Dict[str, Any]] = None,
        *,
        query_profile: Optional[Dict[str, Any]] = None,
        entry: Optional[Dict[str, Any]] = None,
    ) -> float:
        profile = query_profile or self._build_query_profile(query)
        query_counter = profile['counter']
        query_terms = profile['query_terms']
        query_total = profile['total']
        normalized_query = profile['normalized_query']

        if entry is not None:
            content_tokens = entry['tokens']
            content_counter = entry['token_counter']
            normalized_content = entry['normalized_content']
            normalized_doc_name = entry['normalized_doc_name']
            normalized_section_title = entry['normalized_section_title']
        else:
            content_tokens = tuple(self._tokenize_for_search(content))
            content_counter = Counter(content_tokens)
            normalized_content = self._normalize_text(content)
            normalized_doc_name = self._normalize_text((metadata or {}).get('doc_name', ''))
            normalized_section_title = self._normalize_text((metadata or {}).get('section_title', ''))

        if not query_counter or not content_tokens:
            return 0.0

        overlap = sum(min(freq, content_counter.get(token, 0)) for token, freq in query_counter.items())
        if overlap <= 0:
            return 0.0

        coverage = overlap / query_total
        precision = overlap / max(1, len(content_tokens))

        exact_boost = 0.18 if normalized_query and normalized_query in normalized_content else 0.0

        matched_terms = sum(1 for token in query_terms if content_counter.get(token, 0) > 0)
        term_match_ratio = matched_terms / max(1, len(query_terms))

        title_boost = 0.08 if normalized_query and normalized_doc_name and normalized_query in normalized_doc_name else 0.0
        section_boost = 0.06 if normalized_query and normalized_section_title and normalized_query in normalized_section_title else 0.0

        score = (
            0.55 * coverage
            + 0.20 * min(1.0, precision * 8)
            + 0.12 * term_match_ratio
            + exact_boost
            + title_boost
            + section_boost
        )
        return max(0.0, min(1.0, score))

    def _lexical_search(
        self,
        kb_id: str,
        query: str,
        top_k: int,
        *,
        query_expansion_enabled: bool = True,
    ) -> List[Dict[str, Any]]:
        query_profile = self._build_query_profile(query, query_expansion_enabled=query_expansion_enabled)
        search_data = self._get_kb_search_data(kb_id)
        search_entries = search_data['search_entries']
        inverted_index = search_data.get('inverted_index', {})

        # Use inverted index to collect candidate entry indices
        candidate_indices: set = set()
        for token in query_profile['tokens']:
            posting = inverted_index.get(token)
            if posting:
                candidate_indices.update(posting)

        candidates: List[tuple[float, str, Dict[str, Any]]] = []
        for idx in candidate_indices:
            entry = search_entries[idx]
            chunk = entry['chunk']
            score = self._lexical_score(
                query,
                chunk.content,
                chunk.metadata,
                query_profile=query_profile,
                entry=entry,
            )
            if score <= 0:
                continue
            candidate = {
                'id': chunk.id,
                'content': chunk.content,
                'metadata': chunk.metadata,
                'distance': 1.0 - score,
                '_lexical_score': score,
            }
            if len(candidates) < top_k:
                heapq.heappush(candidates, (score, chunk.id, candidate))
            else:
                heapq.heappushpop(candidates, (score, chunk.id, candidate))

        return [
            item
            for _, __, item in sorted(candidates, key=lambda pair: pair[0], reverse=True)
        ]

    def _content_similarity(
        self,
        left: str,
        right: str,
        *,
        left_tokens: Optional[set] = None,
        right_tokens: Optional[set] = None,
    ) -> float:
        left_tokens = left_tokens if left_tokens is not None else set(_tokenize_fn(left))
        right_tokens = right_tokens if right_tokens is not None else set(_tokenize_fn(right))
        if not left_tokens or not right_tokens:
            return 0.0

        overlap = len(left_tokens & right_tokens)
        if overlap <= 0:
            return 0.0
        return overlap / max(1.0, (len(left_tokens) * len(right_tokens)) ** 0.5)

    def _mmr_select_candidates(
        self,
        candidates: List[Dict[str, Any]],
        top_k: int,
        *,
        lambda_mult: float = 0.78,
    ) -> List[Dict[str, Any]]:
        if not candidates or top_k <= 0:
            return []

        remaining = list(candidates)
        selected: List[Dict[str, Any]] = []
        doc_counts = defaultdict(int)

        while remaining and len(selected) < top_k:
            best_idx = 0
            best_score = float("-inf")

            for idx, candidate in enumerate(remaining):
                relevance = float(
                    candidate.get('_reranked_score', candidate.get('_fusion_score', 0.0)) or 0.0
                )
                novelty_penalty = 0.0
                if selected:
                    novelty_penalty = max(
                        self._content_similarity(
                            candidate.get('content', ''),
                            chosen.get('content', ''),
                            left_tokens=candidate.get('_token_set'),
                            right_tokens=chosen.get('_token_set'),
                        )
                        for chosen in selected
                    )

                doc_id = str((candidate.get('metadata') or {}).get('doc_id', ''))
                doc_penalty = 0.08 * doc_counts.get(doc_id, 0) if doc_id else 0.0
                mmr_score = lambda_mult * relevance - (1.0 - lambda_mult) * novelty_penalty - doc_penalty

                if mmr_score > best_score:
                    best_score = mmr_score
                    best_idx = idx

            chosen = remaining.pop(best_idx)
            selected.append(chosen)
            doc_id = str((chosen.get('metadata') or {}).get('doc_id', ''))
            if doc_id:
                doc_counts[doc_id] += 1

        return selected

    def _hybrid_merge_results(
        self,
        kb_id: str,
        query: str,
        top_k: int,
        vector_results: List[Dict[str, Any]],
        rag_settings: Optional[Dict[str, Any]] = None,
        rerank_endpoint_id: str = "",
    ) -> List[Dict[str, Any]]:
        rag_settings = self._sanitize_rag_settings(rag_settings or {})
        search_data = self._get_kb_search_data(kb_id)
        candidate_top_k = max(int(rag_settings["candidate_top_k"]), top_k)
        lexical_results = self._lexical_search(
            kb_id,
            query,
            top_k=candidate_top_k,
            query_expansion_enabled=bool(rag_settings["query_expansion_enabled"]),
        )
        valid_chunk_ids = set(search_data['by_id'].keys())
        search_entry_map = search_data['entry_map']  # 预构建缓存，无需每次重建
        vector_rank = {
            str(item.get('id', '')): rank
            for rank, item in enumerate(vector_results, 1)
            if str(item.get('id', '')) in valid_chunk_ids
        }
        lexical_rank = {
            str(item.get('id', '')): rank
            for rank, item in enumerate(lexical_results, 1)
            if str(item.get('id', '')) in valid_chunk_ids
        }

        combined: Dict[str, Dict[str, Any]] = {}

        for item in vector_results:
            chunk_id = str(item.get('id', ''))
            if not chunk_id or chunk_id not in valid_chunk_ids:
                continue
            score = max(0.0, 1.0 - float(item.get('distance', 0.0) or 0.0))
            entry = search_entry_map.get(chunk_id)
            combined[chunk_id] = {
                'id': chunk_id,
                'content': item.get('content', '') or (entry['content'] if entry else ''),
                'metadata': dict(item.get('metadata') or (entry['metadata'] if entry else {})),
                '_vector_score': score,
                '_lexical_score': 0.0,
                '_vector_rank': vector_rank.get(chunk_id, 0),
                '_lexical_rank': 0,
                '_token_set': set(entry['tokens']) if entry else set(_tokenize_fn(item.get('content', ''))),
            }

        for item in lexical_results:
            chunk_id = item['id']
            entry = search_entry_map.get(chunk_id)
            record = combined.setdefault(chunk_id, {
                'id': chunk_id,
                'content': item.get('content', '') or (entry['content'] if entry else ''),
                'metadata': dict(item.get('metadata') or (entry['metadata'] if entry else {})),
                '_vector_score': 0.0,
                '_lexical_score': 0.0,
                '_vector_rank': vector_rank.get(chunk_id, 0),
                '_lexical_rank': 0,
                '_token_set': set(entry['tokens']) if entry else set(_tokenize_fn(item.get('content', ''))),
            })
            record['content'] = item.get('content', record['content'])
            record['metadata'] = dict(item.get('metadata') or record['metadata'])
            record['_lexical_score'] = max(record['_lexical_score'], float(item.get('_lexical_score', 0.0) or 0.0))
            record['_lexical_rank'] = lexical_rank.get(chunk_id, 0)

        ranked: List[Dict[str, Any]] = []
        seen_content = set()
        candidate_pool: List[Dict[str, Any]] = []

        # Adaptive fusion weights based on query type
        query_type = classify_query(query)
        weights = get_retrieval_weights(query_type)
        w_lexical = weights["lexical"]
        w_vector = weights["vector"]

        for item in combined.values():
            lexical_score = float(item.get('_lexical_score', 0.0) or 0.0)
            vector_score = float(item.get('_vector_score', 0.0) or 0.0)
            rrf_score = 0.0
            if item.get('_vector_rank'):
                rrf_score += 1.0 / (60.0 + float(item['_vector_rank']))
            if item.get('_lexical_rank'):
                rrf_score += 1.0 / (60.0 + float(item['_lexical_rank']))

            base_score = max(
                lexical_score,
                w_lexical * lexical_score + w_vector * vector_score,
            )
            rrf_component = min(1.0, rrf_score * 70.0)
            fusion_score = base_score + 0.05 * rrf_component * max(lexical_score, vector_score)

            if lexical_score <= 0.0 and fusion_score < 0.18:
                continue
            if fusion_score < 0.06:
                continue

            item['_rrf_score'] = rrf_score
            item['_fusion_score'] = fusion_score
            candidate_pool.append(item)

        kb = self.get_knowledge_base(kb_id)
        window_payloads: List[Dict[str, Any]] = []
        section_payloads: List[Dict[str, Any]] = []
        window_docs: List[str] = []
        section_docs: List[str] = []

        for item in candidate_pool:
            window_payload = self._build_chunk_window(
                kb_id,
                item['id'],
                window=int(rag_settings["window_size"]),
                max_chars=int(rag_settings["window_max_chars"]),
            )
            merged_content = window_payload.get('content') or item.get('content', '')
            section_payload = self._build_section_window(
                kb_id,
                item['id'],
                max_chars=int(rag_settings["section_max_chars"]),
            )
            section_content = section_payload.get('content') or merged_content

            window_payloads.append(window_payload)
            section_payloads.append(section_payload)
            window_docs.append(merged_content)
            section_docs.append(section_content)

        remote_window_scores = self._remote_rerank_documents(
            kb,
            query,
            window_docs,
            rerank_endpoint_id=rerank_endpoint_id,
        )
        remote_section_scores = self._remote_rerank_documents(
            kb,
            query,
            section_docs,
            rerank_endpoint_id=rerank_endpoint_id,
        )

        for idx, item in enumerate(candidate_pool):
            window_payload = window_payloads[idx]
            merged_content = window_docs[idx]
            section_payload = section_payloads[idx]
            section_content = section_docs[idx]

            if remote_window_scores is not None:
                remote_window = dict(remote_window_scores[idx] or {})
                remote_window_score = self._bounded_score(remote_window.get('remote_score', 0.0))
                window_rerank_payload = {
                    **remote_window,
                    'final_score': self._bounded_score(
                        0.84 * remote_window_score
                        + 0.12 * float(item.get('_fusion_score', 0.0) or 0.0)
                        + 0.04 * float(item.get('_lexical_score', 0.0) or 0.0)
                    ),
                }
            else:
                window_rerank_payload = self.reranker.score(
                    query=query,
                    snippet=item.get('content', ''),
                    content=merged_content,
                    metadata=dict(item.get('metadata') or {}),
                    retrieval_prior=float(item.get('_fusion_score', 0.0) or 0.0),
                    lexical_score=float(item.get('_lexical_score', 0.0) or 0.0),
                    vector_score=float(item.get('_vector_score', 0.0) or 0.0),
                )
                window_rerank_payload['source'] = 'local_reranker'

            if remote_section_scores is not None:
                remote_section = dict(remote_section_scores[idx] or {})
                remote_section_score = self._bounded_score(remote_section.get('remote_score', 0.0))
                section_rerank_payload = {
                    **remote_section,
                    'final_score': self._bounded_score(
                        0.84 * remote_section_score
                        + 0.12 * float(item.get('_fusion_score', 0.0) or 0.0)
                        + 0.04 * float(item.get('_lexical_score', 0.0) or 0.0)
                    ),
                }
            else:
                section_rerank_payload = self.reranker.score(
                    query=query,
                    snippet=item.get('content', ''),
                    content=section_content,
                    metadata={
                        **dict(item.get('metadata') or {}),
                        'section_title': section_payload.get('section_title') or (item.get('metadata') or {}).get('section_title', ''),
                    },
                    retrieval_prior=float(item.get('_fusion_score', 0.0) or 0.0),
                    lexical_score=float(item.get('_lexical_score', 0.0) or 0.0),
                    vector_score=float(item.get('_vector_score', 0.0) or 0.0),
                )
                section_rerank_payload['source'] = 'local_reranker'

            section_score = float(section_rerank_payload.get('final_score', 0.0) or 0.0)
            window_score = float(window_rerank_payload.get('final_score', 0.0) or 0.0)
            section_chunk_count = len(section_payload.get('chunk_ids') or [])
            window_chunk_count = len(window_payload.get('chunk_ids') or [])
            use_section = (
                section_score > window_score + 0.015
                or (
                    section_chunk_count > window_chunk_count
                    and section_score >= window_score - 0.01
                )
            )
            chosen_payload = section_payload if use_section else window_payload
            chosen_content = section_content if use_section else merged_content
            chosen_rerank_payload = section_rerank_payload if use_section else window_rerank_payload
            reranked_score = float(chosen_rerank_payload.get('final_score', 0.0) or 0.0)

            item['_window_payload'] = window_payload
            item['_section_payload'] = section_payload
            item['_context_payload'] = chosen_payload
            item['_merged_content'] = chosen_content
            item['_rerank_payload'] = chosen_rerank_payload
            item['_window_rerank_payload'] = window_rerank_payload
            item['_section_rerank_payload'] = section_rerank_payload
            item['_context_mode'] = 'section' if use_section else 'window'
            item['_reranked_score'] = reranked_score

        ordered = sorted(
            candidate_pool,
            key=lambda current: float(current.get('_reranked_score', current.get('_fusion_score', 0.0)) or 0.0),
            reverse=True,
        )
        selected_candidates = self._mmr_select_candidates(
            ordered[:max(candidate_top_k, top_k * 2)],
            top_k=max(top_k * 2, top_k),
            lambda_mult=float(rag_settings["mmr_lambda"]),
        )

        for item in selected_candidates:
            normalized_content = self._normalize_text(item.get('content', ''))[:240]
            if not normalized_content or normalized_content in seen_content:
                continue

            lexical_score = float(item.get('_lexical_score', 0.0) or 0.0)
            vector_score = float(item.get('_vector_score', 0.0) or 0.0)
            score = float(item.get('_fusion_score', 0.0) or 0.0)
            reranked_score = float(item.get('_reranked_score', score) or score)
            window_payload = dict(item.get('_window_payload') or {})
            context_payload = dict(item.get('_context_payload') or window_payload)
            section_payload = dict(item.get('_section_payload') or {})
            merged_content = item.get('_merged_content') or context_payload.get('content') or item.get('content', '')
            merged_ids = context_payload.get('chunk_ids') or [item['id']]
            rerank_payload = dict(item.get('_rerank_payload') or {})
            context_mode = str(item.get('_context_mode', 'window') or 'window')

            if reranked_score < float(rag_settings["retrieval_score_threshold"]):
                continue

            result = {
                'id': item['id'],
                'content': merged_content,
                'snippet': item.get('content', ''),
                'metadata': {
                    **dict(item.get('metadata') or {}),
                    'retrieved_chunk_id': item['id'],
                    'merged_chunk_ids': merged_ids,
                    'window_start': window_payload.get('window_start'),
                    'window_end': window_payload.get('window_end'),
                    'section_chunk_ids': section_payload.get('chunk_ids') or [],
                    'section_start': section_payload.get('section_start'),
                    'section_end': section_payload.get('section_end'),
                    'section_title': section_payload.get('section_title') or (item.get('metadata') or {}).get('section_title'),
                    'context_mode': context_mode,
                    'rerank_features': rerank_payload,
                },
                'distance': max(0.0, 1.0 - reranked_score),
                'score': round(reranked_score, 6),
                'retrieval_score': round(reranked_score, 6),
                'hybrid_score': round(score, 6),
                'reranker_score': round(reranked_score, 6),
                'lexical_score': round(lexical_score, 6),
                'vector_score': round(vector_score, 6),
                'rrf_score': round(float(item.get('_rrf_score', 0.0) or 0.0), 6),
            }
            ranked.append(result)
            seen_content.add(normalized_content)

            if len(ranked) >= top_k:
                break

        return ranked

    def search(
        self,
        kb_id: str,
        query: str,
        top_k: int = 5,
        *,
        embedding_endpoint_id: str = "",
        rerank_endpoint_id: str = "",
    ) -> List[Dict[str, Any]]:
        """Search knowledge base with hybrid lexical + vector retrieval."""
        query = (query or '').strip()
        if not query:
            return []

        kb = self.get_knowledge_base(kb_id)
        rag_settings = self._get_rag_settings(kb)
        try:
            top_k = max(1, int(top_k or rag_settings["search_top_k"]))
        except (TypeError, ValueError):
            top_k = int(rag_settings["search_top_k"])

        cached_results = self._get_cached_query_results(
            kb_id,
            query,
            top_k,
            embedding_endpoint_id=embedding_endpoint_id,
            rerank_endpoint_id=rerank_endpoint_id,
        )
        if cached_results is not None:
            return cached_results

        query_embedding = None
        try:
            query_embedding = self._embed_query_for_kb(
                kb,
                query,
                embedding_endpoint_id=embedding_endpoint_id,
            )
        except Exception as exc:
            print(f"[knowledge_base] query embedding fallback: {exc}")
            query_embedding = None

        try:
            vector_results = self.vector_store.search(
                kb_id,
                query,
                top_k=max(int(rag_settings["candidate_top_k"]), top_k),
                query_embedding=query_embedding,
            )
        except Exception:
            vector_results = []

        results = self._hybrid_merge_results(
            kb_id,
            query,
            top_k,
            vector_results,
            rag_settings=rag_settings,
            rerank_endpoint_id=rerank_endpoint_id,
        )
        self._cache_query_results(
            kb_id,
            query,
            top_k,
            results,
            embedding_endpoint_id=embedding_endpoint_id,
            rerank_endpoint_id=rerank_endpoint_id,
        )
        return self._clone_search_results(results)

    def _save_knowledge_base(self, kb: KnowledgeBase):
        """保存知识库"""
        kb_file = os.path.join(self.kb_dir, f"{kb.id}.json")
        with open(kb_file, 'w', encoding='utf-8') as f:
            json.dump({
                'id': kb.id,
                'name': kb.name,
                'description': kb.description,
                'documents': kb.documents,
                'metadata': kb.metadata,
                'created_at': kb.created_at,
                'updated_at': kb.updated_at
            }, f, ensure_ascii=False, indent=2)

    def _save_document(self, doc: Document):
        """保存文档"""
        doc_file = os.path.join(self.doc_dir, f"{doc.id}.json")
        with open(doc_file, 'w', encoding='utf-8') as f:
            json.dump({
                'id': doc.id,
                'name': doc.name,
                'type': doc.type,
                'size': doc.size,
                'content': doc.content,
                'chunks': [
                    {
                        'id': chunk.id,
                        'content': chunk.content,
                        'metadata': chunk.metadata,
                        'search_data': chunk.search_data,
                    }
                    for chunk in doc.chunks
                ],
                'metadata': doc.metadata,
                'created_at': doc.created_at
            }, f, ensure_ascii=False, indent=2)


# 全局实例
knowledge_manager = KnowledgeBaseManager()
